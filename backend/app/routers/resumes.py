import io
import uuid
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update
from app.database import get_db
from app.models import Resume, Profile
from app.schemas.resume import ResumeCreate, ResumeUpdate, ResumeResponse, ResumeParseRequest, ResumeParseResponse
from app.middleware.auth import get_current_user
from app.services.resume_parser import parse_resume_text, extract_skills_from_data
from app.services.gemini import parse_resume_with_ai
from app.services.matching import recompute_scores_for_resume


def extract_text_from_file(content: bytes, file_type: str | None) -> str:
    if file_type == "pdf":
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)
    if file_type == "docx":
        import docx
        document = docx.Document(io.BytesIO(content))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    return content.decode("utf-8", errors="replace")

router = APIRouter(prefix="/api/v1/resumes", tags=["resumes"])


@router.get("", response_model=list[ResumeResponse])
async def list_resumes(
    current_user: Profile = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Resume).where(Resume.user_id == current_user.id).order_by(Resume.created_at.desc())
    )
    return result.scalars().all()


@router.get("/current", response_model=ResumeResponse | None)
async def get_current_resume(
    current_user: Profile = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Resume).where(Resume.user_id == current_user.id, Resume.is_current == True)
    )
    return result.scalar_one_or_none()


@router.get("/{resume_id}", response_model=ResumeResponse)
async def get_resume(
    resume_id: uuid.UUID,
    current_user: Profile = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Resume).where(Resume.id == resume_id, Resume.user_id == current_user.id)
    )
    resume = result.scalar_one_or_none()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return resume


@router.post("/parse", response_model=ResumeParseResponse)
async def parse_resume_text_endpoint(
    data: ResumeParseRequest,
    current_user: Profile = Depends(get_current_user),
):
    raw_text = data.raw_text
    parsed_data = parse_resume_text(raw_text)
    skills = extract_skills_from_data(parsed_data)

    try:
        ai_parsed = await parse_resume_with_ai(raw_text)
        if ai_parsed and ai_parsed.get("contact"):
            parsed_data = ai_parsed
            skills = ai_parsed.get("skills", skills)
    except Exception:
        pass

    return ResumeParseResponse(parsed_data=parsed_data, skills=skills)


@router.post("/upload", response_model=ResumeResponse)
async def upload_resume(
    file: UploadFile = File(...),
    current_user: Profile = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if file.size and file.size > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 5MB)")

    content = await file.read()

    file_type = None
    if file.filename:
        ext = file.filename.rsplit(".", 1)[-1].lower()
        if ext in ("pdf", "docx", "txt"):
            file_type = ext

    raw_text = extract_text_from_file(content, file_type)

    parsed_data = parse_resume_text(raw_text)
    skills = extract_skills_from_data(parsed_data)

    try:
        ai_parsed = await parse_resume_with_ai(raw_text)
        if ai_parsed and ai_parsed.get("contact"):
            parsed_data = ai_parsed
            skills = ai_parsed.get("skills", skills)
    except Exception:
        pass

    await db.execute(
        update(Resume).where(Resume.user_id == current_user.id, Resume.is_current == True).values(is_current=False)
    )

    version_result = await db.execute(
        select(Resume.version).where(Resume.user_id == current_user.id).order_by(Resume.version.desc()).limit(1)
    )
    last_version = version_result.scalar() or 0

    resume = Resume(
        user_id=current_user.id,
        file_name=file.filename or "resume",
        file_type=file_type,
        parsed_data=parsed_data,
        raw_text=raw_text,
        skills=skills,
        version=last_version + 1,
        is_current=True,
    )
    db.add(resume)
    await db.flush()
    await db.refresh(resume)
    await recompute_scores_for_resume(db, resume.id)
    return resume


@router.post("", response_model=ResumeResponse)
async def create_resume_manual(
    data: ResumeCreate,
    current_user: Profile = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await db.execute(
        update(Resume).where(Resume.user_id == current_user.id, Resume.is_current == True).values(is_current=False)
    )

    version_result = await db.execute(
        select(Resume.version).where(Resume.user_id == current_user.id).order_by(Resume.version.desc()).limit(1)
    )
    last_version = version_result.scalar() or 0

    resume = Resume(
        user_id=current_user.id,
        file_name=data.file_name,
        file_type=data.file_type,
        parsed_data=data.parsed_data,
        raw_text=data.raw_text,
        skills=data.skills,
        version=last_version + 1,
        is_current=data.is_current,
    )
    db.add(resume)
    await db.flush()
    await db.refresh(resume)
    await recompute_scores_for_resume(db, resume.id)
    return resume


@router.put("/{resume_id}", response_model=ResumeResponse)
async def update_resume(
    resume_id: uuid.UUID,
    data: ResumeUpdate,
    current_user: Profile = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Resume).where(Resume.id == resume_id, Resume.user_id == current_user.id)
    )
    resume = result.scalar_one_or_none()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    for key, value in data.model_dump(exclude_unset=True).items():
        setattr(resume, key, value)
    await db.flush()
    await db.refresh(resume)
    await recompute_scores_for_resume(db, resume.id)
    return resume


@router.delete("/{resume_id}")
async def delete_resume(
    resume_id: uuid.UUID,
    current_user: Profile = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Resume).where(Resume.id == resume_id, Resume.user_id == current_user.id)
    )
    resume = result.scalar_one_or_none()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    await db.delete(resume)
    return {"detail": "Deleted"}


@router.post("/{resume_id}/restore", response_model=ResumeResponse)
async def restore_resume(
    resume_id: uuid.UUID,
    current_user: Profile = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Restore a previous resume version by creating a new current version from it."""
    result = await db.execute(
        select(Resume).where(Resume.id == resume_id, Resume.user_id == current_user.id)
    )
    resume = result.scalar_one_or_none()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # Unset current for all user's resumes
    await db.execute(
        update(Resume).where(Resume.user_id == current_user.id, Resume.is_current == True).values(is_current=False)
    )

    # Get next version number
    version_result = await db.execute(
        select(Resume.version).where(Resume.user_id == current_user.id).order_by(Resume.version.desc()).limit(1)
    )
    last_version = version_result.scalar() or 0

    # Create new resume from the old one
    new_resume = Resume(
        user_id=current_user.id,
        file_name=resume.file_name,
        file_type=resume.file_type,
        parsed_data=resume.parsed_data,
        raw_text=resume.raw_text,
        skills=resume.skills,
        version=last_version + 1,
        is_current=True,
    )
    db.add(new_resume)
    await db.flush()
    await db.refresh(new_resume)
    await recompute_scores_for_resume(db, new_resume.id)
    return new_resume
